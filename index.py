import random
import os
import sys
import tempfile
import shutil
from collections import defaultdict
from datetime import datetime

# Optional Pillow
try:
    from PIL import Image, ImageTk
    PIL_AVAILABLE = True
except Exception:
    Image = None
    ImageTk = None
    PIL_AVAILABLE = False

# GUI imports
try:
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox, scrolledtext
    try:
        from tkinterdnd2 import DND_FILES, TkinterDnD
        DND_AVAILABLE = True
    except ImportError:
        TkinterDnD = tk
        DND_AVAILABLE = False
except ImportError:
    tk = None

# Optional DOCX support
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except Exception:
    Document = None
    Inches = None
    DOCX_AVAILABLE = False

# Optional SVG -> PNG conversion support
try:
    import cairosvg
    CAIROSVG_AVAILABLE = True
except Exception:
    cairosvg = None
    CAIROSVG_AVAILABLE = False


# Helper: convert various formats to PNG usable by python-docx
def _convert_to_png(src_path):
    """Return path to a PNG file converted from src_path, or None if conversion failed/not possible."""
    if not os.path.exists(src_path):
        return None
    base, ext = os.path.splitext(src_path)
    ext = ext.lower()
    # SVG via cairosvg
    if ext == '.svg' and CAIROSVG_AVAILABLE:
        try:
            dst = base + '.png'
            if not os.path.exists(dst):
                cairosvg.svg2png(url=src_path, write_to=dst)
            return dst
        except Exception:
            return None
    # Use Pillow for other conversions (jfif, etc.)
    if PIL_AVAILABLE:
        try:
            dst = base + '_conv.png'
            img = Image.open(src_path)
            # convert to RGB if necessary
            if img.mode in ('RGBA', 'LA'):
                img = img.convert('RGBA')
            else:
                img = img.convert('RGB')
            img.save(dst, format='PNG')
            return dst
        except Exception:
            return None
    return None


def _embed_image_in_paragraph(par, img_path, width_inch=2.5):
    """Try to embed image into a paragraph; try conversions if necessary. Returns True if embedded."""
    try:
        par.add_run().add_picture(img_path, width=Inches(width_inch))
        return True
    except Exception:
        # try conversion
        conv = _convert_to_png(img_path)
        if conv:
            try:
                par.add_run().add_picture(conv, width=Inches(width_inch))
                return True
            except Exception:
                return False
        return False


def _hide_table_borders(table):
    """Remove/hide borders from a docx table for a cleaner inline layout."""
    if not DOCX_AVAILABLE:
        return
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tbl = table._element
        # ensure tblPr exists
        tblPr = getattr(tbl, 'tblPr', None)
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        # remove existing tblBorders if present
        for child in list(tblPr):
            if 'tblBorders' in child.tag:
                tblPr.remove(child)
        tblBorders = OxmlElement('w:tblBorders')
        for b in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            node = OxmlElement(f'w:{b}')
            node.set(qn('w:val'), 'none')
            node.set(qn('w:sz'), '0')
            node.set(qn('w:space'), '0')
            tblBorders.append(node)
        tblPr.append(tblBorders)

        # Also clear borders on each cell
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = getattr(tc, 'tcPr', None)
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                # remove existing tcBorders
                for child in list(tcPr):
                    if 'tcBorders' in child.tag:
                        tcPr.remove(child)
                tcBorders = OxmlElement('w:tcBorders')
                for b in ('top', 'left', 'bottom', 'right'):
                    node = OxmlElement(f'w:{b}')
                    node.set(qn('w:val'), 'none')
                    node.set(qn('w:sz'), '0')
                    node.set(qn('w:space'), '0')
                    tcBorders.append(node)
                tcPr.append(tcBorders)
    except Exception:
        pass

# Global
loaded_questions = []
images_map = {}  # mapping qid -> list of image paths

def read_questions_from_file(filename):
    """Đọc và parse file ngân hàng câu hỏi"""
    import re
    questions = []
    def clean_opt(s):
        s = s.strip()
        # remove surrounding braces or parentheses
        if s.startswith('{') and s.endswith('}'):
            s = s[1:-1].strip()
        s = re.sub(r'^[A-Da-d]\s*[\.|\)]\s*', '', s)
        return s.strip()

    def split_combined_options(s):
        # Split on semicolon, pipe or / and comma but keep order
        parts = re.split(r';|\||,', s)
        return [clean_opt(p) for p in parts if p.strip()]

    try:
        with open(filename, 'r', encoding='utf-8') as f:
            content = f.read().strip()

        lines = [line.strip() for line in content.split('\n') if line.strip()]

        for line in lines:
            # Skip non-question lines except manual trailing ones
            if not line.startswith('Q') and '|' not in line:
                if '(' in line and ')' in line:
                    try:
                        level_part = line.split('(')[1].split(')')[0].strip().upper()
                        question_text = line.split(')')[1].split('(đáp án')[0].strip().rstrip(' –-')
                        qtype = 'MCQ' if any(opt in question_text for opt in ['A.', 'B.', 'C.', 'D.']) else 'ESSAY'
                        questions.append({
                            'id': f"MANUAL{len(questions)+1:03d}",
                            'level': level_part,
                            'type': qtype,
                            'subject': 'Chưa xác định',
                            'question': question_text,
                            'options': ['','','',''] if qtype == 'MCQ' else ['','','',''],
                            'answer': '',
                            'hint': ''
                        })
                    except Exception:
                        continue
                continue

            parts = line.split('|')
            if len(parts) < 5:
                continue

            qid = parts[0].strip()
            level = parts[1].strip().upper()
            qtype_raw = parts[2].strip().upper()
            subject = parts[3].strip()
            question_text = parts[4].strip()

            qtype = 'ESSAY' if 'ESSAY' in qtype_raw or 'TL' in qtype_raw else 'MCQ'

            # tail parts after question_text
            tail = parts[5:]
            opts = ['', '', '', '']
            answer = ''
            hint = ''

            if qtype == 'MCQ':
                # Case 1: explicit four option fields
                if len(tail) >= 4 and any(tail[:4]):
                    o = tail[:4]
                    opts = [clean_opt(x) for x in o]
                    if len(tail) >= 5 and tail[4].strip():
                        answer = tail[4].strip()
                    if len(tail) >= 6:
                        hint = ' '.join(tail[5:]).strip()
                else:
                    # Combined options in one field or options mixed
                    if tail:
                        # if last token is single letter (A-D) treat as answer
                        last = tail[-1].strip()
                        if re.match(r'^[A-Da-d]$', last):
                            answer = last.upper()
                            opt_source = ' | '.join(tail[:-1])
                        else:
                            opt_source = ' | '.join(tail)

                        found_opts = split_combined_options(opt_source)
                        if found_opts:
                            for i in range(min(4, len(found_opts))):
                                opts[i] = found_opts[i]
                        # try to detect answer inside tail if not set
                        if not answer:
                            for t in tail:
                                if re.match(r'^[A-Da-d]$', t.strip()):
                                    answer = t.strip().upper()
                                    break
                                # or last token of tail might be like 'B' or 'C'
                            # also check for single-letter at end of line
                            m = re.search(r'\|?\s*([A-Da-d])\s*$', line)
                            if m:
                                answer = m.group(1).upper()
                        # everything else after options perhaps hint
                        if len(found_opts) < len(split_combined_options(opt_source)):
                            hint = ''
                        # if still empty, try remaining tail parts as hint
                        if not hint and len(tail) > (1 if answer else 0):
                            possible = tail[-1] if not answer else ' '.join(tail[1:])
                            hint = possible.strip()
            else:
                # Essay: join any tail as hint/notes
                hint = ' '.join(tail).strip()

            questions.append({
                'id': qid,
                'level': level,
                'type': qtype,
                'subject': subject,
                'question': question_text,
                'options': opts,
                'answer': answer,
                'hint': hint
            })
    except Exception as e:
        try:
            messagebox.showerror("Lỗi đọc file", f"Không thể đọc file:\n{e}")
        except Exception:
            print(f"Lỗi đọc file: {e}")
    return questions


def format_question(q, index, show_answers=False, show_hints=False):
    """Định dạng đẹp một câu hỏi"""
    lines = [f"Câu {index}: [{q['level']}] {q['question']}"]
    if q['type'] == 'MCQ' and any(q['options']):
        for i, opt in enumerate(q['options']):
            if opt.strip():
                lines.append(f"  {chr(65 + i)}. {opt}")
        if show_answers and q.get('answer'):
            lines.append(f"  Đáp án: {q['answer']}")
    if show_hints and q.get('hint'):
        lines.append(f"  Gợi ý: {q['hint']}")
    return "\n".join(lines)


def generate_versions(questions, N, req_level, req_mcq, req_essay):
    """Tạo N mã đề theo yêu cầu mức độ và số lượng TN/TL"""
    bank = {'NB': {'MCQ': [], 'ESSAY': []},
            'TH': {'MCQ': [], 'ESSAY': []},
            'VD': {'MCQ': [], 'ESSAY': []},
            'VDH': {'MCQ': [], 'ESSAY': []}}
    
    for q in questions:
        lev = q['level']
        if lev not in bank:
            lev = 'NB'
        typ = q['type'] if q['type'] in ['MCQ', 'ESSAY'] else 'ESSAY'
        bank[lev][typ].append(q)
    
    versions = []
    for _ in range(N):
        selected = []
        remain_mcq = req_mcq
        remain_essay = req_essay
        
        # Phân bổ theo mức độ trước
        for lev in ['NB', 'TH', 'VD', 'VDH']:
            need = req_level.get(lev, 0)
            if need == 0:
                continue
            mcq_pool = bank[lev]['MCQ'].copy()
            essay_pool = bank[lev]['ESSAY'].copy()
            random.shuffle(mcq_pool)
            random.shuffle(essay_pool)
            
            take_mcq = min(need, remain_mcq, len(mcq_pool))
            take_essay = min(need - take_mcq, remain_essay, len(essay_pool))
            
            selected.extend(mcq_pool[:take_mcq])
            selected.extend(essay_pool[:take_essay])
            
            remain_mcq -= take_mcq
            remain_essay -= take_essay
        
        # Bổ sung thêm nếu còn thiếu
        all_mcq = [q for lev in bank for q in bank[lev]['MCQ']]
        all_essay = [q for lev in bank for q in bank[lev]['ESSAY']]
        random.shuffle(all_mcq)
        random.shuffle(all_essay)
        
        selected.extend(all_mcq[:remain_mcq])
        selected.extend(all_essay[:remain_essay])
        
        random.shuffle(selected)
        versions.append(selected)
    
    return versions


def _write_exam_file(path, exam, idx, source_filename=None, copied_map=None, as_docx=False):
    """Write one exam to path. Supports TXT and DOCX (if as_docx=True).
    If as_docx=True, `Document` must be available and copied_map (qid->list of filenames)
    will be used to embed images where possible.
    """
    if as_docx:
        if not DOCX_AVAILABLE:
            print("python-docx không được cài đặt; không thể lưu .docx")
            return False
        try:
            doc = Document()
            doc.add_heading(f"MÃ ĐỀ {idx}", level=1)
            doc.add_paragraph(f"Tạo lúc {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            # ensure MCQ appear before ESSAY in the output
            ordered = [q for q in exam if q.get('type') == 'MCQ'] + [q for q in exam if q.get('type') == 'ESSAY']
            for q_idx, q in enumerate(ordered, 1):
                # create a two-column table: left for question text, right for attachments (images/files)
                tbl = doc.add_table(rows=1, cols=2)
                tbl.autofit = True
                left_cell = tbl.rows[0].cells[0]
                right_cell = tbl.rows[0].cells[1]

                # left: question and options
                lp = left_cell.paragraphs[0]
                lp.add_run(f"Câu {q_idx}: [{q['level']}] ").bold = True
                lp.add_run(q['question'])
                if q['type'] == 'MCQ' and any(q['options']):
                    for i, opt in enumerate(q['options']):
                        if opt.strip():
                            left_cell.add_paragraph(f"{chr(65+i)}. {opt}")
                if q.get('answer'):
                    left_cell.add_paragraph(f"Đáp án: {q.get('answer')}")
                if q.get('hint'):
                    left_cell.add_paragraph(f"Gợi ý: {q.get('hint')}")

                # right: embed attachments (images inlined). If multiple, stack vertically
                for dst in (copied_map or {}).get(q.get('id'), []) or []:
                    dst_path = os.path.join(os.path.dirname(path), 'images', dst)
                    ext = os.path.splitext(dst)[1].lower()
                    embedded = False
                    # SVG: try convert then embed
                    if ext == '.svg':
                        conv = _convert_to_png(dst_path)
                        if conv:
                            rp = right_cell.add_paragraph()
                            embedded = _embed_image_in_paragraph(rp, conv, width_inch=2.5)
                    # raster types (including .jfif)
                    elif ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.jfif']:
                        rp = right_cell.add_paragraph()
                        embedded = _embed_image_in_paragraph(rp, dst_path, width_inch=2.5)

                    if not embedded:
                        right_cell.add_paragraph(f"Tệp đính kèm: images/{dst}")
                # add spacing
                doc.add_paragraph('')
            doc.save(path)
            return True
        except Exception as e:
            print(f"Lỗi khi ghi docx {path}: {e}")
            return False
    else:
        try:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(f"MÃ ĐỀ {idx} — Tạo lúc {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                ordered = [q for q in exam if q.get('type') == 'MCQ'] + [q for q in exam if q.get('type') == 'ESSAY']
                for q_idx, q in enumerate(ordered, 1):
                    f.write(format_question(q, q_idx, show_answers=True, show_hints=True) + "\n\n")
            return True
        except Exception as e:
            print(f"Lỗi khi ghi file {path}: {e}")
            return False


def save_exams_to_directory(exams, source_filename=None, target_dir=None, as_docx=False):
    """Save each exam to a separate file in target_dir. Choose DOCX when as_docx=True."""
    if target_dir is None:
        target_dir = os.getcwd()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = os.path.splitext(os.path.basename(source_filename))[0] if source_filename else f"de_thi_{timestamp}"
    saved_files = []
    images_folder = os.path.join(target_dir, 'images')
    os.makedirs(images_folder, exist_ok=True)
    for i, exam in enumerate(exams, 1):
        suffix = chr(64 + i) if i <= 26 else str(i)
        ext = '.docx' if as_docx else '.txt'
        fn = f"{base_name}_de_{suffix}_{timestamp}{ext}"
        path = os.path.join(target_dir, fn)
        # Copy attachments used in this exam into images_folder and build mapping original->copied_basename
        copied_map = {}
        for q in exam:
            for im in images_map.get(q.get('id'), []) or []:
                if os.path.exists(im):
                    dst_name = os.path.basename(im)
                    dst_path = os.path.join(images_folder, dst_name)
                    # avoid overwrite by renaming
                    if os.path.exists(dst_path):
                        basef, extn = os.path.splitext(dst_name)
                        k = 1
                        while os.path.exists(os.path.join(images_folder, f"{basef}_{k}{extn}")):
                            k += 1
                        dst_name = f"{basef}_{k}{extn}"
                        dst_path = os.path.join(images_folder, dst_name)
                    try:
                        shutil.copy2(im, dst_path)
                        copied_map.setdefault(q.get('id'), []).append(dst_name)
                        copied_map.setdefault(im, dst_name)
                    except Exception as e:
                        print(f"Không thể copy tệp {im}: {e}")
        ok = _write_exam_file(path, exam, suffix, source_filename, copied_map=copied_map, as_docx=as_docx)
        # append attachment references at end of file (for txt files)
        if copied_map and not as_docx:
            try:
                with open(path, 'a', encoding='utf-8') as f:
                    f.write('\nTỆP ĐÍNH KÈM:\n')
                    for q in exam:
                        for dst in copied_map.get(q.get('id'), []):
                            f.write(f"{q.get('id')}: images/{dst}\n")
            except Exception as e:
                print(f"Không thể ghi thông tin tệp đính kèm: {e}")
        if ok:
            saved_files.append(path)
    return saved_files


def save_exams_combined(exams, filepath, source_filename=None):
    try:
        out_dir = os.path.dirname(filepath) or os.getcwd()
        images_folder = os.path.join(out_dir, 'images')
        os.makedirs(images_folder, exist_ok=True)
        # copy all attachments used across exams and record mapping original->copied basename
        copied_map = {}
        for exam in exams:
            for q in exam:
                for im in images_map.get(q.get('id'), []) or []:
                    if os.path.exists(im):
                        dst_name = os.path.basename(im)
                        dst_path = os.path.join(images_folder, dst_name)
                        if os.path.exists(dst_path):
                            base, ext = os.path.splitext(dst_name)
                            k = 1
                            while os.path.exists(os.path.join(images_folder, f"{base}_{k}{ext}")):
                                k += 1
                            dst_name = f"{base}_{k}{ext}"
                            dst_path = os.path.join(images_folder, dst_name)
                        try:
                            shutil.copy2(im, dst_path)
                            copied_map.setdefault(q.get('id'), []).append(dst_name)
                            copied_map.setdefault(im, dst_name)
                        except Exception as e:
                            print(f"Không thể copy tệp {im}: {e}")

        # If filepath indicates .docx and DOCX_AVAILABLE then write docx combined
        if filepath.lower().endswith('.docx'):
            if not DOCX_AVAILABLE:
                print("python-docx không được cài đặt; không thể lưu .docx")
                return None
            try:
                doc = Document()
                doc.add_heading('KẾT QUẢ TẠO ĐỀ THI', level=1)
                doc.add_paragraph(f"Thời gian: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph(f"Số mã đề: {len(exams)}")
                for idx, exam in enumerate(exams, 1):
                    doc.add_page_break()
                    doc.add_heading(f"MÃ ĐỀ {idx}", level=2)
                    ordered = [q for q in exam if q.get('type') == 'MCQ'] + [q for q in exam if q.get('type') == 'ESSAY']
                    for q_idx, q in enumerate(ordered, 1):
                        # create a two-column table so image(s) appear beside question
                        tbl = doc.add_table(rows=1, cols=2)
                        tbl.autofit = True
                        _hide_table_borders(tbl)
                        left = tbl.rows[0].cells[0]
                        right = tbl.rows[0].cells[1]
                        left.add_paragraph(f"Câu {q_idx}: [{q['level']}] {q['question']}")
                        if q['type'] == 'MCQ' and any(q['options']):
                            for i, opt in enumerate(q['options']):
                                if opt.strip():
                                    left.add_paragraph(f"{chr(65+i)}. {opt}")
                        if q.get('answer'):
                            left.add_paragraph(f"Đáp án: {q.get('answer')}")

                        for dst in copied_map.get(q.get('id'), []) or []:
                            dst_path = os.path.join(images_folder, dst)
                            ext = os.path.splitext(dst)[1].lower()
                            embedded = False
                            if ext == '.svg':
                                conv = _convert_to_png(dst_path)
                                if conv:
                                    rp = right.add_paragraph()
                                    embedded = _embed_image_in_paragraph(rp, conv, width_inch=4)
                            elif ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.jfif']:
                                rp = right.add_paragraph()
                                embedded = _embed_image_in_paragraph(rp, dst_path, width_inch=4)
                            if not embedded:
                                right.add_paragraph(f"Tệp đính kèm: images/{dst}")
                        doc.add_paragraph('')
                doc.save(filepath)
                return filepath
            except Exception as e:
                print(f"Lỗi khi lưu docx: {e}")
                return None

        # default: write plain text
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"KẾT QUẢ TẠO ĐỀ THI\n")
            f.write(f"Thời gian: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Số mã đề: {len(exams)}\n\n")
            for idx, exam in enumerate(exams, 1):
                f.write(f"{'='*50}\n")
                f.write(f"MÃ ĐỀ {idx}\n")
                f.write(f"{'='*50}\n\n")
                for q_idx, q in enumerate(exam, 1):
                    f.write(format_question(q, q_idx, show_answers=True, show_hints=True) + "\n")
                    # write references using copied_map if available
                    for dst in copied_map.get(q.get('id'), []) or []:
                        f.write(f"Tệp đính kèm: images/{dst}\n")
                    f.write("\n")
                f.write("\n" + "="*50 + "\n\n")
        return filepath
    except Exception as e:
        print(f"Lỗi khi lưu file: {e}")
        return None


def parse_dnd_paths(data):
    paths = []
    data = data.strip()
    if data.startswith('{'):
        cur = ''
        in_brace = False
        for ch in data:
            if ch == '{':
                in_brace = True
                cur = ''
            elif ch == '}':
                in_brace = False
                if cur:
                    paths.append(cur)
                cur = ''
            elif in_brace:
                cur += ch
    else:
        paths = [p.strip() for p in data.split()]
    return paths


class ExamGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Ngân Hàng Câu Hỏi - Tạo Đề Thi Chuyên Nghiệp")
        self.root.geometry("1000x700")
        self.root.attributes('-topmost', True)
        
        style = ttk.Style()
        style.theme_use('clam')  # Theme hiện đại
        
        # Header
        header = ttk.Label(root, text="TẠO ĐỀ THI TỪ NGÂN HÀNG CÂU HỎI", font=("Helvetica", 16, "bold"))
        header.pack(pady=10)
        
        # Control frame
        ctrl_frame = ttk.Frame(root)
        ctrl_frame.pack(fill='x', padx=10, pady=5)
        
        ttk.Button(ctrl_frame, text="Chọn File", command=self.browse_file).pack(side='left', padx=5)
        ttk.Button(ctrl_frame, text="Upload Hình", command=self.upload_images).pack(side='left', padx=5)
        self.status_label = ttk.Label(ctrl_frame, text="Chưa tải file nào", foreground="gray")
        self.status_label.pack(side='left', padx=20)
        
        # Drop zone
        drop_frame = ttk.LabelFrame(root, text="Kéo & Thả File Vào Đây")
        drop_frame.pack(fill='x', padx=10, pady=10)
        self.drop_label = ttk.Label(drop_frame, text="Thả file .txt vào đây để tải ngân hàng câu hỏi", 
                                   background="#f0f0f0", padding=20)
        self.drop_label.pack(fill='both', expand=True)
        
        if DND_AVAILABLE:
            self.drop_label.drop_target_register(DND_FILES)
            self.drop_label.dnd_bind('<<Drop>>', self.on_drop)
        else:
            ttk.Label(drop_frame, text="Drag & Drop không khả dụng (cài tkinterdnd2 để bật)", foreground="red").pack()
        
                # Config frame
        config_frame = ttk.LabelFrame(root, text="Cấu Hình Đề Thi")     
        config_frame.pack(fill='x', padx=10, pady=5)
        
        row1 = ttk.Frame(config_frame)
        row1.pack(fill='x', padx=10, pady=5)
        ttk.Label(row1, text="Số mã đề:").pack(side='left')
        self.num_versions = tk.IntVar(value=2)
        ttk.Spinbox(row1, from_=1, to=50, width=5, textvariable=self.num_versions).pack(side='left', padx=5)
        
        ttk.Label(row1, text="   Trắc nghiệm:").pack(side='left', padx=10)
        self.num_mcq = tk.IntVar(value=10)
        ttk.Spinbox(row1, from_=0, to=100, width=5, textvariable=self.num_mcq).pack(side='left', padx=5)
        
        ttk.Label(row1, text="   Tự luận:").pack(side='left', padx=10)
        self.num_essay = tk.IntVar(value=4)
        ttk.Spinbox(row1, from_=0, to=50, width=5, textvariable=self.num_essay).pack(side='left', padx=5)
        
        # Level config
        row2 = ttk.Frame(config_frame)
        row2.pack(fill='x', padx=10, pady=5)
        levels = ['NB', 'TH', 'VD', 'VDH']
        self.level_vars = {}
        for i, lev in enumerate(levels):
            ttk.Label(row2, text=f"{lev}:").grid(row=0, column=i*2, padx=5)
            var = tk.IntVar(value=[5,4,3,2][i])
            ttk.Spinbox(row2, from_=0, to=50, width=5, textvariable=var).grid(row=0, column=i*2+1, padx=5)
            self.level_vars[lev] = var
        # Buttons
        btn_frame = ttk.Frame(root)
        btn_frame.pack(fill='x', padx=10, pady=10)
        ttk.Button(btn_frame, text="Tạo Mã Đề", command=self.generate_exams).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="Lưu Đề Ra File", command=self.save_exams).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="Quản lý hình", command=self.manage_images).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="Đóng", command=root.destroy).pack(side='right', padx=5)
        
        # Display
        self.text_display = scrolledtext.ScrolledText(root, wrap='word', font=("Consolas", 11))
        self.text_display.pack(fill='both', expand=True, padx=10, pady=10)
        
        self.root.bind('<Control-o>', lambda e: self.browse_file())

    def update_status(self, text):
        self.status_label.config(text=text)

    def browse_file(self):
        fp = filedialog.askopenfilename(
            title="Chọn file ngân hàng câu hỏi",
            filetypes=[("Text files", "*.txt *.dat *.inp"), ("All files", "*.*")]
        )
        if fp:
            self.load_file(fp)

    def on_drop(self, event):
        paths = parse_dnd_paths(event.data)
        for p in paths:
            p = p.strip('"{}')
            if os.path.isfile(p):
                self.load_file(p)
                break

    def upload_images(self):
        if not loaded_questions:
            messagebox.showwarning("Chưa có ngân hàng", "Vui lòng tải file ngân hàng câu hỏi trước khi upload hình.")
            return
        fps = filedialog.askopenfilenames(title='Chọn tệp đính kèm (được hỗ trợ mọi loại tệp)', filetypes=[('All files','*.*')])
        if not fps:
            return
        # Try to auto-detect QID from filename
        qids = [q['id'] for q in loaded_questions]
        auto_matches = {}
        for fp in fps:
            name = os.path.basename(fp)
            found = None
            for qid in qids:
                if qid.lower() in name.lower():
                    found = qid
                    break
            if found:
                auto_matches[fp] = found

        # Open dialog to assign
        dlg = tk.Toplevel(self.root)
        dlg.title('Gán tệp đính kèm cho câu hỏi')
        dlg.geometry('780x480')
        frm = ttk.Frame(dlg, padding=10)
        frm.pack(fill='both', expand=True)

        ttk.Label(frm, text='Gõ mã câu (QID) hoặc danh sách QID cách nhau bởi dấu phẩy. Để trống để bỏ qua.', wraplength=700).pack(pady=(0,8))

        canvas = tk.Canvas(frm)
        scrollbar = ttk.Scrollbar(frm, orient='vertical', command=canvas.yview)
        inner = ttk.Frame(canvas)
        inner.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox('all')))
        canvas.create_window((0,0), window=inner, anchor='nw')
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        # Try to import PIL for thumbnails
        try:
            from PIL import Image, ImageTk
            PIL_AVAILABLE = True
        except Exception:
            PIL_AVAILABLE = False

        rows = []
        for fp in fps:
            base = os.path.basename(fp)
            row = ttk.Frame(inner)
            row.pack(fill='x', pady=6, padx=4)
            left = ttk.Frame(row)
            left.pack(side='left')
            # show thumbnail only for image file types
            ext = os.path.splitext(base)[1].lower()
            if ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp'] and PIL_AVAILABLE:
                try:
                    img = Image.open(fp)
                    img.thumbnail((80,80))
                    tkimg = ImageTk.PhotoImage(img)
                    lbl = ttk.Label(left, image=tkimg)
                    lbl.image = tkimg
                    lbl.pack()
                except Exception:
                    ttk.Label(left, text='[Hình không xem được]').pack()
            else:
                ttk.Label(left, text='[Tệp]').pack()

            mid = ttk.Frame(row)
            mid.pack(side='left', fill='x', expand=True, padx=8)
            ttk.Label(mid, text=base).pack(anchor='w')
            var = tk.StringVar()
            # prefill with auto match if available
            if fp in auto_matches:
                var.set(auto_matches[fp])
            entry = ttk.Entry(mid, textvariable=var, width=30)
            entry.pack(anchor='w', pady=2)
            ttk.Label(mid, text='Ví dụ: Q001 hoặc Q001,Q002').pack(anchor='w')
            rows.append((fp, var))

        def apply_assign():
            assigned = 0
            skipped = []
            qid_map = {q.upper(): q for q in qids}
            for fp, var in rows:
                txt = var.get().strip()
                if not txt:
                    continue
                # support multiple QIDs separated by commas
                for token in [t.strip() for t in txt.split(',') if t.strip()]:
                    t = token.upper()
                    if t in qid_map:
                        real = qid_map[t]
                        if fp not in images_map.setdefault(real, []):
                            images_map.setdefault(real, []).append(fp)
                            assigned += 1
                    else:
                        # try partial match (token contained in a known QID)
                        candidates = [v for k, v in qid_map.items() if t in k]
                        if len(candidates) == 1:
                            real = candidates[0]
                            if fp not in images_map.setdefault(real, []):
                                images_map.setdefault(real, []).append(fp)
                                assigned += 1
                        else:
                            skipped.append(token)
            dlg.destroy()
            msg = f'Đã gán {assigned} tệp vào câu tương ứng.'
            if skipped:
                msg += f"\nBỏ qua {len(skipped)} mã không tìm thấy: {', '.join(skipped)}"
            messagebox.showinfo('Xong', msg)

        ttk.Button(frm, text='Áp dụng', command=apply_assign).pack(pady=8)

    def manage_images(self):
        # simple viewer of images_map
        dlg = tk.Toplevel(self.root)
        dlg.title('Quản lý hình ảnh gán')
        dlg.geometry('600x400')
        frm = ttk.Frame(dlg, padding=10)
        frm.pack(fill='both', expand=True)
        text = scrolledtext.ScrolledText(frm, wrap=tk.WORD)
        text.pack(fill='both', expand=True)
        if not images_map:
            text.insert('1.0', 'Chưa có hình được gán.')
        else:
            for qid, fps in images_map.items():
                text.insert(tk.END, f"{qid}:\n")
                for p in fps:
                    text.insert(tk.END, f"  - {p}\n")
                text.insert(tk.END, "\n")
        text.config(state=tk.DISABLED)

    def load_file(self, path):
        global loaded_questions
        questions = read_questions_from_file(path)
        if not questions:
            messagebox.showwarning("Không có dữ liệu", "Không tìm thấy câu hỏi hợp lệ trong file.")
            return
        
        loaded_questions = questions
        self.current_filename = path
        total = len(questions)
        mcq_count = sum(1 for q in questions if q['type'] == 'MCQ')
        essay_count = total - mcq_count
        
        self.text_display.delete('1.0', tk.END)
        self.text_display.insert(tk.END, f"ĐÃ TẢI THÀNH CÔNG FILE:\n{os.path.basename(path)}\n\n")
        self.text_display.insert(tk.END, f"Tổng cộng: {total} câu hỏi\n")
        self.text_display.insert(tk.END, f"  • Trắc nghiệm: {mcq_count} câu\n")
        self.text_display.insert(tk.END, f"  • Tự luận: {essay_count} câu\n\n")
        self.text_display.insert(tk.END, "Xem trước 15 câu đầu:\n" + "-"*50 + "\n")
        for i, q in enumerate(questions[:15], 1):
            self.text_display.insert(tk.END, format_question(q, i) + "\n\n")
        if total > 15:
            self.text_display.insert(tk.END, f"... và {total-15} câu nữa.\n")
        
        self.update_status(f"Đã tải {total} câu hỏi")

    def generate_exams(self):
        if not loaded_questions:
            messagebox.showwarning("Chưa có dữ liệu", "Vui lòng tải file ngân hàng câu hỏi trước!")
            return
        
        N = self.num_versions.get()
        req_mcq = self.num_mcq.get()
        req_essay = self.num_essay.get()
        req_level = {lev: self.level_vars[lev].get() for lev in ['NB','TH','VD','VDH']}
        total_per_exam = req_mcq + req_essay
        
        # Thông báo cấu hình
        info = f"SẮP TẠO {N} MÃ ĐỀ\n"
        info += f"Mỗi đề có: {total_per_exam} câu\n"
        info += f"  • Trắc nghiệm: {req_mcq} câu\n"
        info += f"  • Tự luận: {req_essay} câu\n"
        if any(req_level.values()):
            info += "Phân bổ theo mức độ:\n"
            for lev, cnt in req_level.items():
                if cnt > 0:
                    info += f"    - {lev}: {cnt} câu\n"
        
        if not messagebox.askyesno("Xác nhận tạo đề", info + "\nTiếp tục?"):
            return
        
        versions = generate_versions(loaded_questions, N, req_level, req_mcq, req_essay)
        self.generated_versions = versions
        
        self.text_display.delete('1.0', tk.END)
        self.text_display.insert(tk.END, f"HOÀN THÀNH! ĐÃ TẠO THÀNH CÔNG {N} MÃ ĐỀ\n")
        self.text_display.insert(tk.END, f"Mỗi đề có khoảng {total_per_exam} câu (có thể thay đổi nhẹ nếu thiếu câu)\n\n")
        self.text_display.insert(tk.END, "="*80 + "\n\n")
        
        for i, ver in enumerate(versions, 1):
            actual_total = len(ver)
            actual_mcq = sum(1 for q in ver if q['type'] == 'MCQ')
            self.text_display.insert(tk.END, f"MÃ ĐỀ {i} ({actual_total} câu: {actual_mcq} TN + {actual_total-actual_mcq} TL)\n")
            self.text_display.insert(tk.END, "-"*60 + "\n")
            ordered = [q for q in ver if q.get('type') == 'MCQ'] + [q for q in ver if q.get('type') == 'ESSAY']
            for j, q in enumerate(ordered, 1):
                    self.text_display.insert(tk.END, format_question(q, j) + "\n")
                    # show associated files
                    for im in images_map.get(q.get('id'), []) or []:
                        self.text_display.insert(tk.END, f"Tệp đính kèm: {os.path.basename(im)}\n")
                    self.text_display.insert(tk.END, "\n")
            self.text_display.insert(tk.END, "\n" + "="*80 + "\n\n")
        
        self.update_status(f"Đã tạo {N} mã đề")

    def save_exams(self):
        if not self.text_display.get('1.0', 'end').strip():
            messagebox.showinfo("Chưa có nội dung", "Chưa tạo đề để lưu!")
            return
        # If we have generated versions, offer per-file save option
        if getattr(self, 'generated_versions', None):
            sep = messagebox.askyesno("Lưu mỗi mã đề 1 file?", "Bạn có muốn lưu mỗi mã đề vào 1 file riêng không?")
            if sep:
                # ask for target dir
                target_dir = filedialog.askdirectory(title="Chọn thư mục lưu các mã đề")
                if not target_dir:
                    return
                # ask whether to save as Word .docx
                as_docx = messagebox.askyesno("Định dạng lưu", "Bạn có muốn lưu mỗi mã đề thành file Word (.docx) thay vì .txt?")
                if as_docx and not DOCX_AVAILABLE:
                    messagebox.showwarning("Thiếu thư viện", "Không thể lưu .docx vì python-docx chưa được cài. Cài bằng: pip install python-docx")
                    as_docx = False
                saved = save_exams_to_directory(self.generated_versions, getattr(self, 'current_filename', None), target_dir, as_docx=as_docx)
                if saved:
                    messagebox.showinfo("Đã lưu", f"Đã lưu {len(saved)} file vào:\n{target_dir}")
                return

        # default: save combined file
        fp = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text file", "*.txt"), ("All files", "*.*")],
            title="Lưu đề thi"
        )
        if fp:
            success = None
            if getattr(self, 'generated_versions', None):
                # write combined from generated versions (avoid duplicate preview content differences)
                res = save_exams_combined(self.generated_versions, fp, getattr(self, 'current_filename', None))
                success = bool(res)
                if not success and fp.lower().endswith('.docx') and not DOCX_AVAILABLE:
                    messagebox.showwarning("Thiếu thư viện", "Không thể lưu .docx vì python-docx chưa được cài. Cài bằng: pip install python-docx")
            else:
                try:
                    with open(fp, 'w', encoding='utf-8') as f:
                        f.write(self.text_display.get('1.0', tk.END))
                    success = True
                except Exception as e:
                    messagebox.showerror("Lỗi", f"Không thể lưu file: {e}")
                    success = False
            if success:
                messagebox.showinfo("Thành công", f"Đã lưu vào:\n{fp}")
            else:
                # no success, notify
                messagebox.showerror("Thất bại", "Lưu file không thành công.")


def main():
    if tk is None:
        print("Tkinter không khả dụng. Chạy ở chế độ CLI không được hỗ trợ trong phiên bản này.")
        return
    
    root = TkinterDnD.Tk() if DND_AVAILABLE else tk.Tk()
    app = ExamGeneratorApp(root)
    root.mainloop()


if __name__ == "__main__":
    random.seed(42)  # Để tái tạo kết quả
    main()